from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "output" / "docx" / "mapa-societario-documentacion-registro-propiedad-intelectual.docx"
FLOW_PNG = ROOT / "tmp" / "registro-rpi-render" / "mapa-societario-diagrama-flujo.png"

DATE_ES = "28 de agosto de 2026"
COMMIT_FULL = "539776235c996d9953d13256470552c9fd294169"
COMMIT_SHORT = "5397762"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20262E"
MUTED = "5D6772"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FB"
WHITE = "FFFFFF"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA} DXA: {widths_dxa}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_keep_with_next(paragraph, enabled=True) -> None:
    paragraph.paragraph_format.keep_with_next = enabled


def set_repeat_on_each_page(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def add_numbering_definitions(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abstract_start = max(existing_abstract or [0]) + 1
    num_start = max(existing_num or [0]) + 1

    def create(abstract_id: int, num_id: int, fmt: str, text: str, font: str | None = None) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, lvl_jc, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)

    create(abstract_start, num_start, "bullet", "•", "Symbol")
    create(abstract_start + 1, num_start + 1, "decimal", "%1.")
    return num_start, num_start + 1


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_bullet(doc, text: str, num_id: int, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    apply_num(p, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        set_run_font(p.add_run(text))


def add_numbered(doc, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    apply_num(p, num_id)
    set_run_font(p.add_run(text))


def add_body(doc, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        a = p.add_run(bold_lead)
        set_run_font(a, bold=True)
        b = p.add_run(text[len(bold_lead):])
        set_run_font(b)
    else:
        set_run_font(p.add_run(text))


def add_callout(doc, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    set_run_font(r, color=NAVY, bold=True)
    set_run_font(p.add_run(text), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_on_each_page(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), size=10, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run_font(p.add_run(value), size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section_heading(doc, text: str, level=1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def configure_section(section) -> None:
    # Named filing override to the standard_business_brief preset:
    # A4 with 22.5 mm margins, appropriate for a Spanish administrative filing.
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(22.5)
    section.bottom_margin = Mm(22.5)
    section.left_margin = Mm(22.5)
    section.right_margin = Mm(22.5)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.text = "MAPA SOCIETARIO  ·  DOCUMENTACIÓN TÉCNICA"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D7DBE2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    set_run_font(p.add_run(f"Versión documentada · {DATE_ES} · Página "), size=8.5, color=MUTED)
    add_field(p, "PAGE")


def find_font(size: int, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            try:
                return ImageFont.truetype(candidate, size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def multiline(draw, box, text, font, fill, align="center", spacing=10):
    x0, y0, x1, y1 = box
    max_width = x1 - x0 - 54
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total = sum(heights) + spacing * max(0, len(lines) - 1)
    y = y0 + ((y1 - y0) - total) / 2
    for line, height in zip(lines, heights):
        width = draw.textbbox((0, 0), line, font=font)[2]
        x = x0 + ((x1 - x0) - width) / 2 if align == "center" else x0 + 28
        draw.text((x, y), line, font=font, fill=fill)
        y += height + spacing


def arrow(draw, x, y1, y2, color=(46, 116, 181), width=8):
    draw.line((x, y1, x, y2 - 22), fill=color, width=width)
    draw.polygon([(x, y2), (x - 17, y2 - 28), (x + 17, y2 - 28)], fill=color)


def generate_flowchart(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    W, H = 2100, 2550
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    title_font = find_font(60, bold=True)
    box_font = find_font(37, bold=True)
    small_font = find_font(31, bold=False)
    label_font = find_font(27, bold=True)

    draw.text((W / 2, 85), "Flujo principal de funcionamiento", font=title_font, fill=(23, 54, 93), anchor="ma")
    draw.text((W / 2, 155), "Mapa Societario - búsqueda, tratamiento y salida", font=small_font, fill=(93, 103, 114), anchor="ma")

    x0, x1 = 280, 1820
    box_h, gap = 180, 92
    y = 255
    boxes = [
        ("1. Entrada del usuario", "Nombre de empresa o persona; selección de una sugerencia", (232, 238, 245), (23, 54, 93)),
        ("2. Validación y resolución", "Normalización del término y vinculación con la entidad registral", (244, 247, 251), (32, 38, 46)),
        ("3. Consulta de servicios", "API de datos societarios y, cuando procede, funciones de borde", (232, 238, 245), (23, 54, 93)),
        ("4. Tratamiento de datos", "Caché, reintentos, deduplicación, estados de cargo y relaciones", (244, 247, 251), (32, 38, 46)),
        ("5. Construcción del modelo", "Conversión a nodos y enlaces; asociación de metadatos y evidencias", (232, 238, 245), (23, 54, 93)),
        ("6. Presentación interactiva", "Grafo, ficha, tabla, filtros, cronología y rutas de conexión", (244, 247, 251), (32, 38, 46)),
        ("7. Acción o resultado", "Ampliar la red, guardar/exportar, monitorizar o generar informes", (232, 238, 245), (23, 54, 93)),
    ]
    centers = []
    for idx, (heading, detail, fill, text_color) in enumerate(boxes):
        y0, y1 = y, y + box_h
        draw.rounded_rectangle((x0, y0, x1, y1), radius=28, fill=fill, outline=(180, 192, 205), width=4)
        draw.rounded_rectangle((x0 + 24, y0 + 26, x0 + 500, y1 - 26), radius=18, fill=(46, 116, 181))
        multiline(draw, (x0 + 24, y0 + 26, x0 + 500, y1 - 26), heading, box_font, (255, 255, 255))
        multiline(draw, (x0 + 530, y0 + 20, x1 - 20, y1 - 20), detail, small_font, text_color, align="left")
        centers.append((W // 2, y0, y1))
        y = y1 + gap

    for idx in range(len(centers) - 1):
        arrow(draw, W // 2, centers[idx][2] + 8, centers[idx + 1][1] - 8)

    # Decision/loop annotation beside the main flow.
    loop_x = 1870
    draw.line((loop_x, centers[5][1] + 80, loop_x, centers[2][1] + 80), fill=(93, 103, 114), width=5)
    draw.polygon([(loop_x, centers[2][1] + 55), (loop_x - 13, centers[2][1] + 82), (loop_x + 13, centers[2][1] + 82)], fill=(93, 103, 114))
    draw.line((centers[5][0] + 770, centers[5][1] + 80, loop_x, centers[5][1] + 80), fill=(93, 103, 114), width=5)
    draw.text((loop_x - 10, (centers[5][1] + centers[2][1]) / 2), "AMPLIAR\nOTRO NODO", font=label_font, fill=(93, 103, 114), anchor="mm", stroke_width=8, stroke_fill="white")

    draw.text((W / 2, H - 85), "Los resultados interactivos pueden alimentar una nueva consulta o producir una salida conservable.", font=small_font, fill=(93, 103, 114), anchor="ma")
    img.save(path, format="PNG", optimize=True)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MAPA SOCIETARIO")
    set_run_font(r, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_run_font(p.add_run("Documentación técnica del programa de ordenador"), size=16, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    set_run_font(p.add_run("Memoria descriptiva · Diagrama de flujo · Lenguajes · Entorno de ejecución"), size=11.5, color=MUTED)

    table = doc.add_table(rows=5, cols=2)
    set_table_geometry(table, [2700, 6660])
    cover_rows = [
        ("Destino", "Registro Territorial de la Propiedad Intelectual de Madrid"),
        ("Tipo de obra", "Programa de ordenador"),
        ("Autor indicado en el código", "Alessandro Nurnberg"),
        ("Versión documentada", "Aplicación Android 1.29.0 · paquete web 0.1.0"),
        ("Fecha de referencia", DATE_ES),
    ]
    for i, (label, value) in enumerate(cover_rows):
        left, right = table.rows[i].cells
        set_cell_shading(left, LIGHT_BLUE)
        set_cell_shading(right, WHITE)
        set_run_font(left.paragraphs[0].add_run(label), size=10, color=NAVY, bold=True)
        set_run_font(right.paragraphs[0].add_run(value), size=10.5)
        left.paragraphs[0].paragraph_format.space_after = Pt(0)
        right.paragraphs[0].paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(56)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Identificación técnica del estado depositado"), size=9.5, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"Repositorio Git · revisión {COMMIT_SHORT}"), size=9.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("Documento elaborado a partir del código fuente del proyecto."), size=9.5, color=MUTED, italic=True)


def build_document() -> None:
    generate_flowchart(FLOW_PNG)
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    bullet_id, decimal_id = add_numbering_definitions(doc)

    core = doc.core_properties
    core.title = "Mapa Societario - Documentación técnica para el Registro de la Propiedad Intelectual"
    core.subject = "Memoria descriptiva, diagrama de flujo, lenguajes y entorno de ejecución"
    core.author = "Alessandro Nurnberg"
    core.keywords = "Mapa Societario, programa de ordenador, memoria descriptiva, diagrama de flujo"
    core.comments = "Estado técnico de referencia: commit " + COMMIT_FULL

    add_cover(doc)
    doc.add_page_break()

    add_section_heading(doc, "1. Objeto, identificación y alcance", 1)
    add_body(doc, "El presente documento describe técnicamente el programa de ordenador denominado Mapa Societario a efectos de acompañar su presentación ante el Registro Territorial de la Propiedad Intelectual de Madrid. La memoria se refiere al estado del código fuente identificado en este documento y no a una versión futura o distinta del programa.")
    add_callout(doc, "Objeto", "Aplicación de inteligencia de relaciones societarias que permite buscar empresas y personas vinculadas con sociedades españolas, representar sus conexiones en un grafo interactivo y producir vistas, cronologías, instantáneas e informes derivados.")

    add_section_heading(doc, "1.1 Identificación del programa", 2)
    add_table(doc, ["Elemento", "Identificación"], [
        ["Denominación", "Mapa Societario"],
        ["Categoría", "Aplicación web, aplicación híbrida Android y extensión de navegador, con funciones serverless auxiliares"],
        ["Autor indicado en los encabezados del código", "Alessandro Nurnberg"],
        ["Versión de la aplicación Android", "1.29.0 (versionCode 36)"],
        ["Versión de los paquetes web", "0.1.0"],
        ["Estado de referencia", f"Commit Git {COMMIT_FULL}, de {DATE_ES}"],
        ["Identificador Android", "es.mapasocietario.app"],
        ["Dominio de explotación web", "mapasocietario.es"],
    ], [2700, 6660])

    add_section_heading(doc, "1.2 Alcance de la documentación", 2)
    add_body(doc, "Se documentan el núcleo web interactivo, los servicios de acceso y transformación de datos, las funciones ejecutadas en el borde, la persistencia local, la envoltura Android y la extensión de Chrome incluidas en el repositorio. Los servicios externos consumidos por la aplicación se describen como dependencias de ejecución y no como parte del código objeto de esta memoria.")

    add_section_heading(doc, "2. Memoria descriptiva de la aplicación", 1)
    add_section_heading(doc, "2.1 Finalidad y problema resuelto", 2)
    add_body(doc, "Mapa Societario facilita la exploración de relaciones entre sociedades españolas, sus administradores, apoderados, cargos y socios únicos. El programa transforma resultados procedentes de fuentes y servicios registrales en una representación gráfica navegable, con herramientas para distinguir entidades, seguir vínculos, revisar hechos y conservar una investigación.")
    add_body(doc, "La interfaz se ofrece en español e inglés. Su flujo principal está pensado para una consulta sin cuenta: el usuario introduce una denominación social o un nombre personal, selecciona una coincidencia, examina el grafo y amplía las entidades que resulten relevantes.")

    add_section_heading(doc, "2.2 Funciones principales", 2)
    for item in [
        "Búsqueda unificada y autocompletado de empresas, administradores y socios únicos, incluidos nombres anteriores.",
        "Resolución de la identidad societaria y agrupación de resultados vinculados a una misma entidad registral.",
        "Construcción de un grafo de nodos y enlaces para empresas, personas, cargos y relaciones de socio único.",
        "Expansión progresiva de empresas o personas mediante nuevas consultas, sin reconstruir todo el espacio de trabajo.",
        "Filtrado por nombre, notas, vigencia del cargo y categoría; resaltado de conexiones compartidas y cálculo de rutas dentro de la red cargada.",
        "Ficha de empresa, tablas de datos, línea temporal de cargos y visualización de hechos o hallazgos asociados.",
        "Curación local del espacio de trabajo: notas privadas, etiquetas, colores, ocultación, eliminación, fusión y separación de nodos.",
        "Autoguardado local e importación/exportación de instantáneas JSON con nodos, enlaces, vista, filtros y metadatos de investigación.",
        "Generación o acceso a informes de relaciones, informes de due diligence, monitorización y funciones de investigación asistida cuando el usuario las solicita.",
        "Páginas informativas y fichas indexables de empresas, panel estadístico, gestión de pedidos y rutas de alertas." ,
    ]:
        add_bullet(doc, item, bullet_id)

    add_section_heading(doc, "2.3 Componentes funcionales", 2)
    add_table(doc, ["Componente", "Responsabilidad principal"], [
        ["Interfaz React", "Enrutamiento, formularios, navegación bilingüe, estados de carga, fichas, tablas, diálogos y pantallas auxiliares."],
        ["Motor gráfico", "Representación Canvas mediante un grafo de fuerzas; interacción con nodos, enlaces, filtros, cámara y selección."],
        ["Capa de servicios", "Consultas HTTP, reintentos, caché, normalización, resolución de identidad y preparación de datos para la interfaz."],
        ["Persistencia local", "IndexedDB, localStorage y sessionStorage para sesión, preferencias, caché, autoguardado y continuidad de pedidos."],
        ["Funciones Cloudflare Pages", "Renderizado y respuesta de rutas públicas, recepción de feedback y registro controlado de demanda para páginas de empresa."],
        ["Base D1", "Persistencia serverless de señales de demanda, páginas promovidas y métricas operativas separadas del grafo local."],
        ["Aplicación Android", "Empaquetado Capacitor, WebView, navegación nativa, descargas y compras integradas de Google Play."],
        ["Extensión Chrome", "Consulta contextual y panel lateral de solo lectura para mostrar una ficha registral y la red de cargos."],
        ["Worker analítico", "Proceso serverless auxiliar para informes de uso y rendimiento; no interviene en la construcción del grafo del usuario."],
    ], [3000, 6360])

    add_section_heading(doc, "2.4 Arquitectura lógica", 2)
    add_body(doc, "La arquitectura adopta una separación entre presentación, lógica de dominio y servicios externos. La interfaz React mantiene el estado del grafo y delega la obtención de datos a servicios JavaScript. Estos servicios forman las solicitudes, aplican reintentos y caché, normalizan las respuestas y convierten los registros en entidades utilizables por el motor gráfico.")
    add_body(doc, "El sitio web se compila como recursos estáticos con Vite y se publica en Cloudflare Pages. Determinadas rutas incorporan funciones JavaScript en el entorno de borde y una base Cloudflare D1. La aplicación Android reutiliza el mismo paquete web mediante Capacitor y añade integraciones nativas. La extensión Chrome constituye una interfaz complementaria y de solo lectura.")

    add_section_heading(doc, "2.5 Modelo de información y tratamiento", 2)
    add_body(doc, "El modelo principal es un grafo dirigido o no dirigido según la naturaleza de cada relación. Los nodos representan empresas, personas o cargos societarios; los enlaces representan nombramientos, ceses, apoderamientos, participaciones de socio único y otras relaciones derivadas de los datos disponibles. Cada entidad puede conservar metadatos como denominación, estado, fechas, categoría del cargo, referencias y atributos de presentación.")
    add_body(doc, "Antes de presentar los datos, el programa ejecuta tareas de normalización de nombres, resolución por clave estable, deduplicación, clasificación de cargos, cálculo de vigencia y vinculación de enlaces. Las fusiones manuales y notas introducidas por el usuario se mantienen diferenciadas de los datos obtenidos de fuentes externas.")

    doc.add_page_break()
    add_section_heading(doc, "3. Diagrama de flujo", 1)
    add_body(doc, "El siguiente diagrama resume el recorrido funcional principal. La ampliación de una entidad crea un ciclo controlado: se consulta de nuevo la capa de servicios, se incorporan los resultados al modelo existente y se actualiza la visualización sin perder el trabajo previo.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(5)
    picture = p.add_run().add_picture(str(FLOW_PNG), width=Inches(6.22))
    picture._inline.docPr.set("descr", "Diagrama de flujo de Mapa Societario: entrada, validación, consulta, tratamiento, modelo, visualización y resultado, con ciclo de ampliación de nodos.")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)
    set_run_font(caption.add_run("Figura 1. Flujo principal de la aplicación."), size=9, color=MUTED, italic=True)

    add_section_heading(doc, "3.1 Secuencia detallada", 2)
    for item in [
        "El usuario introduce una empresa o persona y selecciona una coincidencia del autocompletado.",
        "La aplicación valida la entrada, normaliza el nombre y, cuando existe, utiliza una clave estable para identificar la entidad.",
        "La capa de servicios consulta las API pertinentes por HTTPS y controla errores, límites y reintentos.",
        "La respuesta se depura, clasifica y deduplica; se calculan estados de cargo y se preservan referencias útiles.",
        "Los registros se transforman en nodos y enlaces, se integran con el grafo ya existente y se recalcula la disposición visual.",
        "El usuario inspecciona, filtra, amplía o edita localmente la vista. Una ampliación vuelve al paso de consulta.",
        "El resultado puede permanecer en autoguardado local, exportarse como instantánea, alimentar un informe o activar servicios opcionales." ,
    ]:
        add_numbered(doc, item, decimal_id)

    add_section_heading(doc, "4. Lenguajes de programación y tecnologías de implementación", 1)
    add_body(doc, "El programa utiliza principalmente JavaScript moderno con módulos ECMAScript. JSX se emplea para declarar la interfaz React. Los demás lenguajes cubren capas nativas, datos, presentación y automatización de construcción.")
    add_table(doc, ["Lenguaje / tecnología", "Uso en la obra", "Ubicación representativa"], [
        ["JavaScript (ES Modules)", "Lógica principal, servicios, transformación de datos, componentes auxiliares, Cloudflare Functions/Workers, pruebas y scripts de construcción.", "src/, functions/, workers/, scripts/"],
        ["JSX", "Componentes React, composición de pantallas, controles del grafo, fichas, diálogos y extensión de navegador.", "src/**/*.jsx, chrome-extension/src/**/*.jsx"],
        ["Java", "Integración nativa Android: actividad principal, descargas y plugin de compras Google Play.", "android/app/src/main/java/"],
        ["SQL", "Definición de tablas, índices y restricciones de la base Cloudflare D1.", "migrations/*.sql"],
        ["HTML", "Documento de entrada, páginas informativas estáticas y panel de la extensión.", "index.html, public/, chrome-extension/src/panel/"],
        ["CSS", "Estilos globales, diseño adaptable y hojas de estilo de recursos web.", "src/index.css, public/vendor/"],
        ["Groovy / Gradle DSL", "Configuración de compilación y dependencias de la aplicación Android.", "android/*.gradle, android/app/*.gradle"],
        ["Python", "Herramienta de mantenimiento para la actualización de datos auxiliares.", "scripts/refresh-registradores.py"],
    ], [2150, 4540, 2670])

    add_section_heading(doc, "4.1 Formatos declarativos y de intercambio", 2)
    add_body(doc, "Se emplean JSON para configuración, datos, manifiestos e instantáneas; XML para recursos y manifiesto Android; TOML para configuración de Cloudflare; CSV para conjuntos auxiliares; y Markdown para documentación técnica. Estos formatos no constituyen el núcleo algorítmico, pero forman parte de la organización y ejecución del programa.")

    add_section_heading(doc, "4.2 Bibliotecas y marcos principales", 2)
    add_table(doc, ["Biblioteca / marco", "Finalidad"], [
        ["React y React DOM", "Modelo de componentes y renderizado de la aplicación web."],
        ["Material UI y Emotion", "Componentes de interfaz, temas y estilos."],
        ["D3, d3-force y react-force-graph-2d", "Cálculo y representación interactiva de redes societarias."],
        ["React Router", "Navegación entre rutas y pantallas."],
        ["Recharts", "Gráficos estadísticos del panel."],
        ["React PDF", "Visualización de documentos PDF en la interfaz."],
        ["Vite y Vitest", "Compilación, servidor de desarrollo y pruebas automatizadas."],
        ["Capacitor", "Empaquetado híbrido y puente con Android."],
        ["Cloudflare Pages, Workers y D1", "Publicación, funciones serverless y persistencia de borde."],
    ], [3100, 6260])

    add_section_heading(doc, "5. Entorno de ejecución", 1)
    add_section_heading(doc, "5.1 Entorno del usuario final", 2)
    add_table(doc, ["Modalidad", "Requisitos y ejecución"], [
        ["Aplicación web", "Navegador moderno con JavaScript y módulos ES habilitados, Canvas 2D, Fetch, almacenamiento web e IndexedDB. Requiere conexión HTTPS para consultar servicios remotos."],
        ["Aplicación Android", "Android 7.0 o posterior (API mínima 24), objetivo API 36. Se ejecuta en una WebView administrada por Capacitor 8 y utiliza APIs nativas para descargas, navegación y compras integradas."],
        ["Extensión Chrome", "Navegador compatible con extensiones Manifest V3, permisos de menú contextual y panel lateral, y acceso HTTPS a api.ncdata.eu."],
    ], [2500, 6860])

    add_section_heading(doc, "5.2 Entorno de servidor y servicios", 2)
    add_table(doc, ["Capa", "Entorno"], [
        ["Sitio web", "Archivos HTML, CSS, JavaScript e imágenes compilados con Vite y publicados en Cloudflare Pages."],
        ["Funciones de borde", "Runtime JavaScript de Cloudflare Pages Functions para rutas públicas, feedback y señales de demanda."],
        ["Persistencia serverless", "Cloudflare D1, base de datos relacional compatible con SQLite, enlazada como SEO_DB o ANALYTICS_DB según el componente."],
        ["Worker analítico", "Cloudflare Workers con compatibilidad Node.js, disparadores programados y acceso a D1."],
        ["Servicios de datos", "API HTTPS de datos societarios en api.ncdata.eu y servicios especializados configurables para pagos e investigación."],
    ], [2650, 6710])

    add_section_heading(doc, "5.3 Entorno de desarrollo, compilación y prueba", 2)
    add_bullet(doc, "Node.js 20 o superior como base común recomendada; la extensión declara expresamente Node.js >= 20.", bullet_id)
    add_bullet(doc, "npm como gestor de dependencias; package-lock.json fija la resolución concreta de paquetes.", bullet_id)
    add_bullet(doc, "Vite 5 para desarrollo y construcción del paquete web; Vitest y node:test para pruebas automatizadas.", bullet_id)
    add_bullet(doc, "Android SDK con compileSdk 36, targetSdk 36 y minSdk 24; Android Gradle Plugin 9.2.1.", bullet_id)
    add_bullet(doc, "Capacitor 8 para sincronizar el paquete web compilado con el proyecto Android.", bullet_id)
    add_bullet(doc, "Wrangler para desarrollo y despliegue de Cloudflare Pages, Workers y enlaces D1.", bullet_id)

    add_callout(doc, "Condición de conectividad", "El grafo puede conservar y reabrir instantáneas locales, pero una búsqueda nueva, la ampliación de entidades, la monitorización, los pagos y la generación remota de informes dependen de servicios HTTPS externos.")

    add_section_heading(doc, "5.4 Variables y configuración", 2)
    add_body(doc, "Las URL de los servicios se centralizan en la configuración de la aplicación y admiten sustitución mediante variables Vite en desarrollo. Las credenciales o secretos de funciones y Workers se mantienen fuera del repositorio y se inyectan en el entorno de despliegue. Los manifiestos de Capacitor, Android, Chrome y Cloudflare definen identificadores, permisos, enlaces de navegación, bases de datos y parámetros de construcción.")

    add_section_heading(doc, "6. Entradas, salidas y persistencia", 1)
    add_table(doc, ["Tipo", "Descripción"], [
        ["Entradas del usuario", "Texto de búsqueda, selección de coincidencias, filtros, acciones sobre nodos, notas, preferencias, ficheros de instantánea y datos de pedido cuando proceda."],
        ["Entradas remotas", "Respuestas JSON de servicios de datos societarios, pagos, alertas, investigación y funciones auxiliares."],
        ["Salidas visuales", "Grafo interactivo, fichas, tablas, líneas temporales, panel estadístico, mensajes de estado y páginas públicas."],
        ["Salidas conservables", "Instantáneas JSON, imágenes de cronologías, texto o tablas copiables y documentos o informes generados por servicios asociados."],
        ["Persistencia local", "IndexedDB para autoguardado y cachés; localStorage/sessionStorage para preferencias, continuidad de sesión y referencias temporales de pedidos."],
        ["Persistencia remota", "D1 para señales y metadatos serverless; almacenamiento y sistemas externos para pedidos, alertas o informes cuando el usuario activa esas funciones."],
    ], [2400, 6960])

    add_section_heading(doc, "6.1 Control de errores y consistencia", 2)
    add_body(doc, "La capa de servicios distingue errores de red, autenticación, ausencia de resultados, límites de frecuencia y fallos de servidor. Las lecturas pueden repetirse con espera exponencial y se cachean para evitar consultas redundantes. Las instantáneas importadas se validan por formato, versión, unicidad de nodos y coherencia de referencias antes de incorporarse al grafo.")
    add_body(doc, "La aplicación mantiene separados los datos recibidos de fuentes externas y las hipótesis o ajustes locales del usuario. Las correcciones, fusiones y notas se muestran como acciones del espacio de trabajo y no modifican por sí mismas el Registro Mercantil ni la publicación oficial.")

    add_section_heading(doc, "7. Elementos propios, dependencias y fuentes externas", 1)
    add_body(doc, "La obra documentada comprende la selección y organización de componentes, la lógica de búsqueda y resolución, las reglas de normalización y vinculación, el modelo del grafo, la interacción, las vistas, la persistencia, las integraciones y el código auxiliar contenido en el repositorio identificado.")
    add_body(doc, "La aplicación incorpora bibliotecas de terceros bajo sus respectivas licencias y consume servicios o datos externos. Esta memoria no atribuye al autor derechos sobre las bibliotecas, plataformas, marcas, API ni datos oficiales de terceros. En particular, los datos del BORME y otros resultados externos sirven como entradas; el programa aporta el tratamiento, la presentación y la interacción descritos.")
    add_callout(doc, "Delimitación", "La protección solicitada para el programa de ordenador debe entenderse referida al código y a la expresión original de la aplicación depositada, sin extenderse a los datos oficiales, servicios externos ni componentes de terceros considerados aisladamente.")

    add_section_heading(doc, "8. Identificación técnica del depósito", 1)
    add_body(doc, "Para vincular esta memoria con una copia concreta del código fuente, se adopta como referencia la revisión Git indicada a continuación. Si el material entregado al Registro se empaqueta en un archivo, se recomienda conservar asimismo el nombre, tamaño y huella SHA-256 del archivo final en el justificante interno del depósito.")
    add_table(doc, ["Dato", "Valor"], [
        ["Repositorio", "mapasocietario"],
        ["Rama de referencia", "Estado de trabajo correspondiente al commit indicado"],
        ["Commit completo", COMMIT_FULL],
        ["Commit abreviado", COMMIT_SHORT],
        ["Fecha del commit", DATE_ES],
        ["Comando de compilación web", "npm run build"],
        ["Comando de pruebas", "npm test"],
        ["Directorio de salida web", "dist/"],
    ], [2600, 6760])

    add_section_heading(doc, "8.1 Archivos representativos", 2)
    for item in [
        "src/main.jsx y src/App.jsx: arranque, rutas y contenedor principal.",
        "src/components/SpanishCompanyNetworkGraph.jsx: interacción y representación del grafo.",
        "src/services/spanishCompaniesService.js: consultas, normalización y lógica de acceso a datos.",
        "src/utils/graphSnapshot.js y src/utils/graphAutosave.js: exportación, validación y persistencia local.",
        "functions/: funciones Cloudflare Pages y páginas server-side auxiliares.",
        "android/: proyecto nativo, configuración Capacitor e integraciones Android.",
        "chrome-extension/: extensión Manifest V3 de consulta contextual.",
        "workers/analytics/: Worker serverless auxiliar de analítica operativa.",
        "package.json, capacitor.config.json y wrangler.toml: dependencias y configuración de ejecución.",
    ]:
        add_bullet(doc, item, bullet_id)

    add_section_heading(doc, "9. Declaración de cierre de la memoria", 1)
    add_body(doc, "La presente memoria describe de forma funcional y técnica el estado de Mapa Societario identificado por la revisión indicada. Sus apartados de diagrama, lenguajes y entorno de ejecución forman parte inseparable de la descripción del programa y permiten comprender su estructura, funcionamiento y requisitos técnicos.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("Autor indicado en el código fuente: Alessandro Nurnberg"), size=10.5, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"Fecha: {DATE_ES}"), size=10.5)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
